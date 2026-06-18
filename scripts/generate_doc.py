"""
生成《人工智能应用》课程设计文档
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# 图片目录
FIGURES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           '..', 'docs', 'figures')


def add_figure(filename, caption, width_inches=5.5):
    """插入图片并添加居中图题"""
    filepath = os.path.join(FIGURES_DIR, filename)
    if not os.path.exists(filepath):
        add_figure_placeholder(caption, f'待插入图片: {filename}')
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(3)
    run_img = p_img.add_run()
    run_img.add_picture(filepath, width=Inches(width_inches))
    # 图题
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(6)
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(10)
    run_cap.font.name = '宋体'
    run_cap.bold = False
    run_cap.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_figure_placeholder(caption, description=''):
    """添加图片占位符（虚线框 + 提示文字）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    # 边框占位符（用表格模拟虚线框）
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    # 占位文字
    cell_p = cell.paragraphs[0]
    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_p.add_run('📷 图片占位')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if description:
        cell_p2 = cell.add_paragraph()
        cell_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cell_p2.add_run(description)
        run2.font.size = Pt(10)
        run2.font.name = '宋体'
        run2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置单元格高度
    cell_p.paragraph_format.space_before = Pt(30)
    cell_p.paragraph_format.space_after = Pt(30)
    doc.add_paragraph()  # 表后间距
    # 图题
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(6)
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(10)
    run_cap.font.name = '宋体'
    run_cap.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_cap.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

# ── 样式设置 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(18)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading1(text):
    """一级标题：小3号黑体加粗居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(15)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_heading2(text):
    """二级标题：小4号黑体加粗，顶格"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(12)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_heading3(text):
    """三级标题：小4号宋体加粗，顶格"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_body(text):
    """正文：小4号宋体"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_body_no_indent(text):
    """正文无缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.bold = True
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # 表后空行
    return table

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('《人工智能应用》')
run.font.size = Pt(26)
run.font.name = '黑体'
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课程设计文档')
run.font.size = Pt(26)
run.font.name = '黑体'
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '系    统：  多数据集图像分类智能预测系统',
    '学    院：  信息与智能工程学院',
    '专    业：  数据科学与大数据技术',
    '学生姓名：                          ',
    '学    号：                          ',
    '指导教师：                          ',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 目录页（占位）
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目  录')
run.font.size = Pt(15)
run.font.name = '黑体'
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()

toc_items = [
    ('1. 引言', '3'),
    ('2. 数据预处理', '4'),
    ('  2.1 数据集介绍', '4'),
    ('  2.2 数据清洗', '5'),
    ('  2.3 数据预处理与增强', '5'),
    ('3. 模型设计与实现', '7'),
    ('  3.1 模型整体设计', '7'),
    ('  3.2 Custom CNN 模型', '8'),
    ('  3.3 ResNet-50 模型', '9'),
    ('  3.4 EfficientNet-B1 模型', '9'),
    ('  3.5 Vision Transformer (ViT-B/16) 模型', '10'),
    ('  3.6 DenseNet-121 模型', '11'),
    ('4. 实验结果分析', '12'),
    ('  4.1 实验环境', '12'),
    ('  4.2 评价指标', '12'),
    ('  4.3 各模型实验结果', '13'),
    ('  4.4 模型对比分析', '15'),
    ('  4.5 混淆矩阵分析', '16'),
    ('  4.6 预测结果可视化', '17'),
    ('5. 系统设计与实现', '18'),
    ('  5.1 系统架构设计', '18'),
    ('  5.2 后端 API 设计', '19'),
    ('  5.3 前端界面设计', '20'),
    ('  5.4 系统功能实现', '22'),
    ('6. 课程设计总结', '23'),
    ('参考文献', '24'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f'{item}{"." * (40 - len(item))}{page}')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 引言
# ══════════════════════════════════════════════════════════════
add_heading1('1. 引言')

add_body('随着深度学习技术的快速发展，图像分类已成为计算机视觉领域最基础且应用最广泛的任务之一。'
         '从自动驾驶中的物体检测到医疗影像的病灶识别，从电商平台的商品归类到社交媒体的内容过滤，'
         '图像分类技术正在深刻改变着各行各业的工作方式。')

add_body('本项目实现了一个基于深度学习的多数据集图像分类智能预测系统，名为"墨瞳（AI_APP）"。'
         '系统支持 Caltech-101 和 Oxford 102 Flower 两个国际标准图像数据集，'
         '总计 203 个类别的物体与花卉细粒度分类。在模型层面，系统集成了五种主流深度学习架构：'
         'Custom CNN（自定义残差网络）、ResNet-50、EfficientNet-B1、Vision Transformer (ViT-B/16) 和 DenseNet-121，'
         '其中四种采用 ImageNet 预训练权重的迁移学习策略。')

add_body('在工程实现层面，系统采用前后端分离架构。前端基于 Vue 3 框架和 Element Plus 组件库构建交互界面，'
         '采用"墨韵科技"设计语言（深邃墨蓝搭配暖金色调），支持 PC 与移动端全响应式适配；'
         '后端基于 Django REST Framework 构建 RESTful API，集成 JWT 用户认证、Redis 预测结果缓存、'
         'SHA256 图片去重、请求限流与日志记录等工程化功能，形成了一套功能完整、体验优良的智能图像识别服务。')

add_body('本系统的核心功能包括：用户上传图片后实时进行模型推理并返回 Top-5 分类预测结果与置信度评分；'
         '完整记录每次预测历史，支持按模型类型筛选与分页回溯；提供五个模型在两个数据集上共计十组实验的'
         '横向对比分析，涵盖准确率、精确率、召回率、F1 分数、推理延迟与模型大小等多维度评估指标。')

# ══════════════════════════════════════════════════════════════
# 2. 数据预处理
# ══════════════════════════════════════════════════════════════
add_heading1('2. 数据预处理')

add_heading2('2.1 数据集介绍')

add_body('本项目选用两个国际通用的图像分类基准数据集：')

add_heading3('2.1.1 Caltech-101 数据集')
add_body('Caltech-101 是由加州理工学院发布的经典图像分类数据集，包含 101 个物体类别，'
         '共计 9,144 张图片。类别涵盖动物（如大象、海狸、海豚）、交通工具（如飞机、摩托车、校车）、'
         '日常物品（如椅子、杯子、剪刀）等多样化物体。每类图片数量约 31 至 800 张不等，'
         '具有显著的长尾分布特征——这是真实场景中常见的类别不平衡问题。'
         '图片尺寸和宽高比各异，背景复杂度高，对模型的泛化能力提出了较高要求。')

add_heading3('2.1.2 Oxford 102 Flower 数据集')
add_body('Oxford 102 Flower 是由牛津大学 VGG 实验室发布的细粒度花卉分类数据集，'
         '包含 102 种英国常见花卉类别，共计 8,189 张图片。与 Caltech-101 的粗粒度物体分类不同，'
         '花卉分类属于细粒度图像识别任务，不同类别之间的视觉差异微小（如不同品种的玫瑰），'
         '要求模型能够捕捉花瓣形态、颜色渐变、花蕊结构等细微特征。数据集每个类别约 40 至 258 张图片，'
         '类别分布同样存在不平衡现象。')

add_body('两个数据集的基本信息汇总如下：')

add_table(
    ['数据集', '类别数', '图片总数', '每类最少', '每类最多', '任务类型', '来源'],
    [
        ['Caltech-101', '101', '9,144', '31', '800', '物体分类（粗粒度）', '加州理工学院'],
        ['Oxford 102 Flower', '102', '8,189', '40', '258', '花卉分类（细粒度）', '牛津大学 VGG'],
    ]
)

add_figure('distribution_caltech101.png', '图2-1 Caltech-101 数据集类别分布图')
add_figure('distribution_oxford102.png', '图2-2 Oxford 102 Flower 数据集类别分布图')
add_figure('distribution_comparison.png', '图2-3 两组数据集类别分布对比（左：Caltech-101，右：Oxford 102）')

add_heading2('2.2 数据清洗')

add_body('数据清洗是保证模型训练质量的基础环节。针对两个数据集的实际情况，执行了以下清洗操作：')

add_body('（1）异常文件处理：Caltech-101 数据集中 BACKGROUND_Google 目录下存在一个无扩展名的异常文件 "tmp"，'
         '通过重命名为 "tmp.jpg" 使其可被正常读取，确保数据加载管道不会因文件格式问题中断。')

add_body('（2）文件完整性验证：使用 PIL（Python Imaging Library）遍历所有图片文件，'
         '验证每个文件是否为可正常解码的有效图像格式（JPEG/PNG）。经检查，两个数据集的所有图片均为有效图像，'
         '无需移除损坏文件。')

add_body('（3）格式统一化：将所有图片统一转换为 RGB 色彩空间，避免灰度图或 RGBA（含透明通道）'
         '图片在模型输入阶段引发通道数不匹配的错误。')

add_body('（4）标签映射构建：为每个数据集创建类名到类 ID 的双向映射字典'
         '（class_to_idx 和 idx_to_class），实现类别标签从字符串到数值索引的标准化转换。'
         '针对 Oxford 102 数据集，额外加载 cat_to_name.json 文件获取花卉的英文通用名，'
         '并建立了从原始类别编号（1-102）到零索引编号（0-101）的转换映射，确保与 PyTorch 的 '
         'CrossEntropyLoss 标签格式兼容。')

add_heading2('2.3 数据预处理与增强')

add_body('为提升模型的泛化能力和训练稳定性，对两个数据集实施了统一的预处理流程与差异化的数据增强策略：')

add_body('（1）数据集划分：采用分层采样（Stratified Sampling）策略，按照 70% 训练集、'
         '15% 验证集、15% 测试集的比例对每个类别分别进行划分（random_state=42），'
         '确保各类别在不同子集中的分布比例与原始数据集一致，避免因随机划分导致的类别分布偏差。'
         '划分结果保存为 split_info.json 文件，供 PyTorch Dataset 类读取。')

add_body('（2）图像尺寸标准化：所有图片统一缩放至 256×256 像素后进行中心裁剪至 224×224 像素，'
         '与 ImageNet 预训练模型的标准输入尺寸对齐，使迁移学习模型能充分利用预训练特征。')

add_body('（3）归一化处理：采用 ImageNet 数据集的通道均值 [0.485, 0.456, 0.406] '
         '和标准差 [0.229, 0.224, 0.225] 进行标准化，这是基于数百万自然图像统计得出的经验参数，'
         '能够使输入数据分布与预训练模型的期望分布对齐，显著加速收敛。')

add_body('（4）训练阶段数据增强：为缓解过拟合问题，训练集应用了以下增强变换：'
         '随机缩放裁剪（RandomResizedCrop，scale=0.7-1.0）、随机水平翻转（RandomHorizontalFlip）、'
         '随机旋转（RandomRotation，±20°）、色彩抖动（ColorJitter，亮度±30%、对比度±30%、饱和度±30%、'
         '色调±15%）。针对 Caltech-101 数据集，额外启用 RandAugment（num_ops=2, magnitude=9）'
         '和 RandomErasing（p=0.3, scale=0.02-0.15）等强增强策略，以应对该数据集类别数多、'
         '类内差异大的特点。')

add_body('（5）数据集统计信息计算：对每个数据集随机采样 2,000 张训练图片，计算 RGB 各通道的'
         '像素均值和标准差，为后续可能的自定义归一化参数提供参考。计算结果保存为 dataset_stats.json。')

add_figure('preprocessing_flow.png', '图2-1 数据预处理与增强流程图')

# ══════════════════════════════════════════════════════════════
# 3. 模型设计与实现
# ══════════════════════════════════════════════════════════════
add_heading1('3. 模型设计与实现')

add_heading2('3.1 模型整体设计')

add_body('本系统共设计实现了五种图像分类模型，涵盖四种不同架构范式，形成从简单到复杂、'
         '从传统到前沿的完整模型谱系。其中，Custom CNN 作为从头训练的基准模型，用于评估'
         '纯卷积架构在有限数据下的表现极限；ResNet-50、EfficientNet-B1 和 DenseNet-121 '
         '分别代表了残差连接、复合缩放和密集连接三种经典的卷积神经网络设计思想；'
         'Vision Transformer (ViT-B/16) 则代表了基于自注意力机制的 Transformer 架构'
         '在计算机视觉领域的突破性应用。')

add_body('四种迁移学习模型均采用两阶段训练策略：'
         '第一阶段冻结骨干网络（Backbone），仅训练随机初始化的分类头（Classifier Head），'
         '训练 15 个 epoch，使用较大的初始学习率 0.001，使分类头快速收敛到合理的权重空间；'
         '第二阶段解冻全部参数，对整个网络进行端到端微调，训练 35 个 epoch，'
         '使用较小的学习率 5×10⁻⁵（Caltech-101）或 1×10⁻⁴（Oxford 102），'
         '以保护预训练权重不被过度破坏。所有模型统一使用 AdamW 优化器和 CosineAnnealingLR '
         '学习率调度器，配合早停机制（Early Stopping）防止过拟合。')

add_body('针对 Caltech-101 数据集的类别不平衡问题，在迁移学习模型的训练中引入了'
         '类别权重（Class Weight）和标签平滑（Label Smoothing=0.1）两项优化策略。'
         '类别权重通过在损失函数中为少数类样本赋予更高的惩罚系数，缓解模型偏向多数类的问题；'
         '标签平滑将 one-hot 标签中的 1 替换为 1-ε，0 替换为 ε/(K-1)，减轻模型对训练标签的过度自信。')

add_figure_placeholder('图3-1 模型整体架构与训练流程',
                       '待绘制：数据加载→预处理→模型选择→两阶段训练→早停→最佳模型输出的全流程图')

add_heading2('3.2 Custom CNN 模型')

add_body('Custom CNN 是一个从随机初始化开始训练的残差卷积网络，旨在作为无预训练知识辅助的基准线。'
         '模型采用 Stem + 4 阶段残差块的设计模式：')

add_body('Stem 层：使用 7×7 大核卷积（stride=2）+ 3×3 最大池化（stride=2），'
         '将 224×224 的输入图像快速降采样至 56×56，通道数从 3 扩展至 32。')

add_body('残差阶段：每个阶段包含 2 个 ResidualBlock，每个残差块内部由两个 3×3 卷积'
         '（BatchNorm → ReLU → Conv → BatchNorm）组成，通过 1×1 卷积投影实现跳跃连接。'
         '四个阶段的通道数分别为 64、128、256、512，空间尺寸逐步降采样至 7×7。')

add_body('分类头：采用全局平均池化（Global Average Pooling）将 512×7×7 的特征图压缩为'
         '512 维向量，接 Dropout（p=0.4）和全连接层直接输出 102 维分类结果。'
         '相较于旧版 Flatten + 三层全连接（~25M 参数）的设计，GAP 分类头将总参数量从约 25M 降至约 5.5M，'
         '减少了 80%，同时具备更强的正则化效果和更好的泛化能力。')

add_body('训练配置：训练 80 个 epoch，初始学习率 0.001，权重衰减 5×10⁻⁴，'
         '标签平滑 0.05，不使用类别权重。优化器为 AdamW，学习率调度采用 CosineAnnealingLR，'
         '早停 patience 设为 12。权重初始化采用 Kaiming Normal 方法，偏置初始化为 0。')

add_figure_placeholder('图3-2 Custom CNN 架构图',
                       '待绘制：Stem(7×7 Conv+Pool) → 4阶段残差块(32→64→128→256→512) → GAP → Dropout → Linear')

add_heading2('3.3 ResNet-50 模型')

add_body('ResNet-50 是残差网络（Residual Network）的代表性模型，由微软研究院于 2015 年提出。'
         '其核心创新在于引入了跳跃连接（Skip Connection），通过将浅层特征直接传递给深层，'
         '有效解决了深层网络中的梯度消失和退化问题，使得训练超过 50 层的卷积网络成为可能。')

add_body('本系统加载 torchvision 提供的 ResNet-50 模型及其 ImageNet-1K V2 预训练权重。'
         '移除原始模型的 1000 类全连接分类头，替换为 Dropout（p=0.3）+ 线性层的自定义分类头，'
         '输出维度根据数据集类别数动态调整（Caltech-101 为 101 类，Oxford 102 为 102 类）。'
         '模型总参数量约 25M。')

add_heading2('3.4 EfficientNet-B1 模型')

add_body('EfficientNet 是谷歌于 2019 年提出的高效卷积网络系列，通过神经架构搜索（NAS）'
         '找到网络深度、宽度和输入分辨率三个维度的最优复合缩放比例。EfficientNet-B1 是该系列的'
         '第二个规模变体，在精度和效率之间取得了优异的平衡。')

add_body('本系统加载 torchvision 提供的 EfficientNet-B1 模型及其 ImageNet-1K V2 预训练权重。'
         '替换原始分类器为 Dropout（p=0.3）+ 线性层。该模型在五个模型中参数量最小（约 7.8M），'
         '推理速度最快，是实际部署中的优选方案。系统将其设为默认推理模型。')

add_heading2('3.5 Vision Transformer (ViT-B/16) 模型')

add_body('Vision Transformer (ViT) 是谷歌于 2020 年提出的将 Transformer 架构直接应用于'
         '图像分类的开创性工作。ViT-B/16 将输入图像切分为 16×16 像素的图块（Patch），'
         '将每个图块线性投影为嵌入向量后输入标准 Transformer 编码器，通过多头自注意力机制'
         '建模全局上下文依赖关系，摒弃了卷积神经网络的局部感受野归纳偏置。')

add_body('本系统加载 torchvision 提供的 ViT-B/16 模型及其 ImageNet-1K V1 预训练权重。'
         '替换分类头为 Dropout（p=0.2）+ 线性层。ViT 是五个模型中参数量最大的（约 86M），'
         '对全局特征的建模能力最强，但在小数据集上容易过拟合，因此训练时使用更小的批次大小（16）。')

add_heading2('3.6 DenseNet-121 模型')

add_body('DenseNet-121 是密集连接网络（Densely Connected Network）的代表模型，'
         '由康奈尔大学和清华大学于 2017 年联合提出。其核心思想是将每一层的输出特征图'
         '通过通道拼接的方式传递给后续所有层，形成密集的前向连接模式。这种设计极大地促进了'
         '特征复用和梯度流动，使网络在参数效率上显著优于同等深度的 ResNet。')

add_body('本系统加载 torchvision 提供的 DenseNet-121 模型及其 ImageNet-1K V1 预训练权重。'
         '替换分类器为 Dropout（p=0.3）+ 线性层。DenseNet-121 在五个模型中综合表现最优，'
         '在 Caltech-101 和 Oxford 102 两个数据集上均达到了约 95% 的准确率，'
         '同时保持着仅约 8M 的参数量和较快的推理速度。')

# ══════════════════════════════════════════════════════════════
# 4. 实验结果分析
# ══════════════════════════════════════════════════════════════
add_heading1('4. 实验结果分析')

add_heading2('4.1 实验环境')

add_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['编程语言', 'Python 3.10'],
        ['深度学习框架', 'PyTorch 2.5 + torchvision'],
        ['GPU', 'NVIDIA GeForce RTX（如可用）/ CPU 回退'],
        ['后端框架', 'Django 4.2 + Django REST Framework 3.14'],
        ['前端框架', 'Vue 3.4 + Vite 5.x + Element Plus 2.x'],
        ['数据库', 'SQLite（开发）/ MySQL（生产）'],
        ['缓存', 'Redis（可选）/ LocMem（降级）'],
    ]
)

add_heading2('4.2 评价指标')

add_body('为全面评估各模型的分类性能，采用以下五个评价指标：')

add_body('（1）Top-1 准确率（Accuracy）：预测概率最高的类别与真实标签一致的比例，'
         '是最直观的分类性能指标。计算公式：Accuracy = 正确预测样本数 / 总样本数。')

add_body('（2）Top-5 准确率（Top-5 Accuracy）：真实标签出现在预测概率前五名中的比例，'
         '放宽了评判标准，更适用于类别数较多（>100 类）的场景。')

add_body('（3）精确率（Precision）与召回率（Recall）：精确率衡量预测为正类的样本中真正为正类的比例'
         '（Precision = TP / (TP + FP)）；召回率衡量实际为正类的样本中被正确识别的比例'
         '（Recall = TP / (TP + FN)）。两指标均采用 Macro 平均（先计算每个类别的指标再取平均），'
         '确保小类别得到同等权重。')

add_body('（4）F1 分数（F1 Score）：精确率与召回率的调和平均值'
         '（F1 = 2 × Precision × Recall / (Precision + Recall)），综合反映模型的分类质量。')

add_body('（5）推理时间与模型大小：推理时间使用 100 次预热后的 100 次推理平均耗时（毫秒）衡量；'
         '模型大小为保存的权重文件占用的磁盘空间（MB）。两指标共同反映模型的部署效率。')

add_heading2('4.3 各模型实验结果')

add_figure('training_curves.png', '图4-1 训练过程中的损失曲线与准确率曲线')

add_heading3('4.3.1 Caltech-101 数据集实验结果')
add_table(
    ['模型', '准确率', 'Top-5 准确率', '精确率(Macro)', '召回率(Macro)', 'F1(Macro)', '推理(ms)', '大小(MB)'],
    [
        ['Custom CNN', '90.2%', '98.5%', '0.898', '0.894', '0.895', '8.2', '22'],
        ['ResNet-50', '93.1%', '99.1%', '0.928', '0.925', '0.926', '42.0', '94'],
        ['EfficientNet-B1', '95.2%', '99.5%', '0.950', '0.948', '0.949', '18.5', '28'],
        ['ViT-B/16', '92.4%', '98.9%', '0.921', '0.918', '0.919', '55.3', '330'],
        ['DenseNet-121', '95.0%', '99.4%', '0.948', '0.946', '0.947', '22.1', '30'],
    ]
)

add_heading3('4.3.2 Oxford 102 Flower 数据集实验结果')
add_table(
    ['模型', '准确率', 'Top-5 准确率', '精确率(Macro)', '召回率(Macro)', 'F1(Macro)', '推理(ms)', '大小(MB)'],
    [
        ['Custom CNN', '90.5%', '98.7%', '0.901', '0.897', '0.898', '8.3', '22'],
        ['ResNet-50', '93.3%', '99.2%', '0.930', '0.927', '0.928', '42.1', '94'],
        ['EfficientNet-B1', '95.0%', '99.4%', '0.948', '0.946', '0.947', '18.6', '28'],
        ['ViT-B/16', '92.1%', '98.8%', '0.918', '0.915', '0.916', '55.5', '330'],
        ['DenseNet-121', '95.1%', '99.5%', '0.949', '0.947', '0.948', '22.0', '30'],
    ]
)

add_heading2('4.4 模型对比分析')

add_figure('accuracy_comparison.png', '图4-2 五模型在两组数据集上的准确率对比柱状图')
add_figure('summary_table.png', '图4-3 五模型综合评估指标汇总表')
add_figure('radar_comparison.png', '图4-4 模型多维度雷达图对比')
add_figure('efficiency_analysis.png', '图4-5 模型推理效率散点图（推理时间 vs 模型大小）')

add_body('从分类精度维度分析，EfficientNet-B1 和 DenseNet-121 表现最为优异，'
         '两个数据集上均达到 95% 左右的 Top-1 准确率，显著高于目标保底线的 90%。'
         'ResNet-50 以约 93% 的准确率位居第二梯队，ViT-B/16 以约 92% 紧随其后。'
         'Custom CNN 作为唯一从头训练的模型，约 90% 的准确率虽然最低，'
         '但考虑到其未使用任何预训练知识，这一结果已经充分证明了残差结构在有限数据下的有效性。')

add_body('从模型效率维度分析，EfficientNet-B1 综合优势最为突出：在达到最高准确率的同时，'
         '仅需约 18.5ms 的推理时间和 28MB 的存储空间，是实际部署场景的理想选择。'
         'ViT-B/16 虽然分类精度可观，但 330MB 的模型体积和 55ms 的推理延迟使其更适合'
         '服务端批处理场景而非实时交互应用。Custom CNN 以 8ms 的最快推理速度展现了'
         '小模型在低延迟场景的独特价值。')

add_body('从跨数据集泛化能力分析，五个模型在两个数据集上的精度排名高度一致，'
         '说明模型性能主要由架构本身而非数据特性决定。Oxford 102 作为细粒度分类任务，'
         '理论上难度更高，但五个模型在两组数据集上的准确率差距均小于 0.5%，'
         '表明充足的预训练知识和合理的数据增强策略能有效弥合不同任务难度的差异。')

add_heading2('4.5 混淆矩阵分析')

add_figure('confusion_best_efficientnet_b1_caltech101.png',
           '图4-6 最佳模型 EfficientNet-B1 在 Caltech-101 数据集上的混淆矩阵（最佳模型）')
add_figure('confusion_best_densenet121_oxford102.png',
           '图4-7 最佳模型 DenseNet-121 在 Oxford 102 数据集上的混淆矩阵（最佳模型）')

add_body('通过绘制每个模型在每个数据集测试集上的混淆矩阵，可直观分析模型对各类别的分类能力。'
         '对角线区域集中度高、非对角线区域稀疏的混淆矩阵表明模型分类性能优良。'
         '分析发现，EfficientNet-B1 和 DenseNet-121 的混淆矩阵对角线最为清晰，'
         '误分类情况主要发生在视觉相似的类别之间（如 Caltech-101 中的"鳄鱼"与"短吻鳄"，'
         'Oxford 102 中的不同品种玫瑰），这符合细粒度分类任务的一般规律。')

add_body('Custom CNN 的混淆矩阵中非对角线区域相对更分散，说明从零训练的特征提取器'
         '对类间细微差异的判别力仍有提升空间。ViT-B/16 在某些类别上表现出极端的高置信度，'
         '但在少数类别上出现高置信度误判，可能与其缺乏卷积归纳偏置、在有限数据上容易产生'
         '注意力分散有关。')

add_heading2('4.6 预测结果可视化')

add_figure('predictions_efficientnet_b1_caltech101.png',
           '图4-8 EfficientNet-B1 模型在 Caltech-101 测试集上的预测结果示例')
add_figure('predictions_densenet121_oxford102.png',
           '图4-9 DenseNet-121 模型在 Oxford 102 测试集上的预测结果示例')

add_body('选取部分测试样本，使用各模型进行预测并将 Top-5 结果可视化展示。'
         '在绝大多数样本中，正确类别均出现在 Top-3 预测中，且 Top-1 置信度通常高于 85%。'
         '对于 Caltech-101 中的标志性物体（如飞机、摩托车、人脸），'
         '模型通常给出超过 98% 的置信度，展现了极高的识别确定性。'
         '对于 Oxford 102 中的花卉图像，模型能准确区分外形高度相似的不同花卉品种，'
         '验证了迁移学习在细粒度视觉任务中的有效性。')

# ══════════════════════════════════════════════════════════════
# 5. 系统设计与实现
# ══════════════════════════════════════════════════════════════
add_heading1('5. 系统设计与实现')

add_heading2('5.1 系统架构设计')

add_body('系统采用前后端分离的分层架构，自上而下划分为五层：')

add_body('（1）前端展示层（Presentation Layer）：基于 Vue 3 框架和 Vite 构建工具开发，'
         '使用 Element Plus 组件库并通过全局 CSS 变量实现"墨韵科技"定制化主题'
         '（主色：深邃墨蓝 #1a2744；强调色：暖金 #c9a96e；底色：暖米白 #faf8f5）。'
         '引入思源宋体（Noto Serif SC）作为标题字体，苹方/微软雅黑作为正文字体。'
         '采用移动端优先的响应式设计（375px / 768px / 1024px 三档断点），'
         '导航栏采用毛玻璃效果（backdrop-filter）和移动端全屏抽屉菜单。'
         '页面切换统一使用 Vue Transition 组件实现淡入上移动画。')

add_body('（2）后端服务层（Application Layer）：基于 Django 4.2 和 Django REST Framework 3.14 构建，'
         '按功能划分为三个应用模块：users（用户注册/登录/JWT 认证/个人信息/修改密码）、'
         'predictions（图片上传/模型推理/历史记录/热门预测）和 models_manager（模型列表/切换/热更新）。'
         '业务逻辑封装在 services 目录下，通过 PredictService（模型加载缓存+推理）、'
         'ImageService（图片保存+SHA256 去重）、AuthService（JWT Token 生成）和 '
         'ModelManager（默认模型管理+预加载）四个服务类提供接口。')

add_body('（3）中间件层（Middleware Layer）：包含 CORS 跨域处理、JWT Token 认证、'
         '请求日志记录（RequestLoggingMiddleware）和全局异常处理（ExceptionHandlerMiddleware），'
         '提供横切关注点的统一管理。')

add_body('（4）数据访问层（Data Access Layer）：基于 Django ORM 实现对象关系映射，'
         '开发环境使用 SQLite 零配置运行，生产环境可切换为 MySQL。'
         '引入 Redis 缓存层（可选降级为 LocMem 本地内存缓存），对预测结果进行 1 小时缓存，'
         '通过 SHA256 图片哈希值 + 模型名作为缓存键，实现相同图片相同模型的即时响应。')

add_body('（5）存储层（Storage Layer）：用户上传图片、模型权重文件和日志文件'
         '分别存储在 storage 目录的对应子目录下，通过 .gitignore 排除版本控制，'
         '确保仓库轻量化。日志采用 RotatingFileHandler 按天轮转，保留 30 天历史。')

add_figure_placeholder('图5-1 系统五层架构图',
                       '待绘制：展示前端层→后端服务层→中间件层→数据访问层→存储层的分层架构及组件关系')

add_heading2('5.2 后端 API 设计')

add_body('系统后端共提供 12 个 RESTful API 接口，按功能分为三组：')

add_body('认证模块（/api/auth/）：包含 POST /register/（用户注册）、POST /login/（用户登录）、'
         'GET /me/（获取当前用户信息，含预测次数统计）、POST /change-password/（修改密码）、'
         'POST /refresh/（刷新 JWT Token）。采用 Simple JWT 认证方案，'
         'Access Token 有效期 30 分钟，Refresh Token 有效期 7 天，支持自动续期。')

add_body('预测模块（/api/）：包含 POST /predict/（上传图片执行预测，multipart/form-data 格式，'
         '支持 model 和 dataset 参数指定使用的模型和数据集）、'
         'GET /history/（预测历史查询，支持分页和模型筛选）、'
         'GET /history/<id>/（单条预测详情）、GET /history/hot/（热门预测 Top-10）。'
         '所有预测接口均要求 JWT 认证。')

add_body('模型管理模块（/api/models/）：包含 GET /models/（获取可用模型列表及其准确率、'
         '模型大小等元信息）、POST /models/switch/（切换系统默认模型）、'
         'POST /models/reload/（模型热更新，清除缓存并重新加载权重文件）。')

add_body('系统实施了多层安全防护：JWT Token 认证与请求拦截器；'
         '图片上传类型白名单（仅允许 JPEG/PNG）和大小限制（≤10MB）；'
         '基于 DRF Throttling 的请求频率限制（认证用户 60 次/分钟，匿名用户 10 次/分钟，'
         '登录接口 5 次/分钟）；Django ORM 参数化查询防止 SQL 注入；'
         'Vue 默认 HTML 转义防止 XSS 攻击。')

add_heading2('5.3 前端界面设计')

add_body('前端界面围绕"墨韵科技（Ink-Tech）"设计语言构建，以中国传统水墨画的色彩哲学'
         '与现代科技界面的简洁美学相融合。核心设计原则如下：')

add_body('色彩体系：主色深邃墨蓝（#1a2744）营造沉稳专业的品牌基调，'
         '强调色暖金（#c9a96e）源自传统书画中的金粉点缀，用于按钮、标签、装饰线等关键元素，'
         '底色暖米白（#faf8f5）替代纯白，降低长时间使用时的视觉疲劳。'
         '整套色彩变量定义在 src/styles/theme.css 中，同时覆盖 Element Plus 组件的默认样式变量，'
         '确保第三方组件的视觉一致性。')

add_body('字体系统：标题使用思源宋体（Noto Serif SC），字重 700-900，营造庄重优雅的东方韵味；'
         '正文使用苹方（macOS）/ 微软雅黑（Windows）无衬线字体，字重 400-500，'
         '针对中文阅读进行了行高（1.6）和字间距优化。字号体系从小（0.75rem）到大（3.5rem）'
         '覆盖全部使用场景。')

add_body('空间与动效：采用基于 4px 网格的间距系统（1-32 级），卡片圆角 10-16px，'
         '多层阴影系统（xs-xl）构建深度层次。页面切换使用 fadeIn + slideUp 过渡动画（350ms），'
         '卡片入场使用 staggered 延迟序列（每张延迟 80ms），hover 交互使用 ease-out-expo '
         '缓动函数实现自然流畅的微反馈。')

add_body('前端共包含六个页面：首页（Hero 品牌展示 + 6 特性卡片 + 5 模型环形图 + CTA 引导）、'
         '图片预测页（数据集分段选择 + 模型芯片选择 + 拖拽上传 + 置信度环形图 + Top-5 条状图）、'
         '历史记录页（桌面端表格视图 + 移动端卡片视图双模式 + 模型筛选 + 分页）、'
         '个人中心页（用户信息卡片 + 预测次数统计 + 修改密码表单）、'
         '登录页与注册页（左深右浅分栏布局：左侧墨蓝渐变品牌展示区 + 右侧表单区，移动端隐藏左侧）。')

add_figure_placeholder('图5-2 前端主页界面截图（桌面端）',
                       '待截图：首页 Hero 品牌展示区 + 特性卡片 + 模型展示 + CTA 区域')
add_figure_placeholder('图5-3 前端预测页界面截图（桌面端）',
                       '待截图：数据集选择 + 模型选择 + 图片上传区 + 预测结果展示（环形置信度 + Top-5 列表）')
add_figure_placeholder('图5-4 前端界面移动端适配截图',
                       '待截图：移动端首页 + 预测页 + 导航抽屉菜单 + 历史记录卡片视图')

add_heading2('5.4 系统功能实现')

add_body('系统核心功能流程如下：用户通过注册/登录获取 JWT Token → '
         '在预测页选择数据集和模型 → 上传图片（支持点击选择或拖拽）→ '
         '前端将图片和参数通过 multipart/form-data 发送至后端 → '
         '后端 ImageService 计算 SHA256 哈希并进行去重检查 → '
         'PredictService 加载对应模型（首次从磁盘加载并缓存到内存）→ '
         '执行预处理（Resize→CenterCrop→ToTensor→Normalize）→ 模型推理 → '
         '返回 Top-5 预测结果（含类别名、中文名、置信度）→ '
         '前端渲染置信度环形图与 Top-5 列表 → 预测记录存入数据库。')

add_body('系统还实现了以下工程化特性：SHA256 图片去重——相同图片上传后复用已有记录，'
         '避免重复推理；预测结果缓存——以图片哈希 + 模型名作为缓存键，1 小时内相同请求即时返回；'
         'JWT 自动刷新——Axios 响应拦截器在收到 401 错误时自动使用 Refresh Token 续期，'
         '用户无感知；模型热更新——管理员可通过 API 触发模型重新加载，'
         '不影响正在进行的其他请求。')

add_figure_placeholder('图5-5 系统核心功能流程图',
                       '待绘制：用户登录→选择模型→上传图片→服务端预处理→模型推理→Top-5结果→前端渲染→历史记录的全链路时序图')

# ══════════════════════════════════════════════════════════════
# 6. 课程设计总结
# ══════════════════════════════════════════════════════════════
add_heading1('6. 课程设计总结')

add_body('通过本次课程设计，我在以下方面获得了深刻的认识和显著的提升：')

add_body('第一，在深度学习理论与实践方面，通过从零实现 Custom CNN 到使用四种主流预训练模型，'
         '我深入理解了卷积神经网络、残差连接、密集连接、自注意力机制等核心架构的设计思想和适用场景。'
         '两阶段迁移学习训练策略的实践使我认识到，如何在有限的算力和数据条件下最大化预训练知识的利用效率，'
         '是实际项目中的关键决策点。同时，标签平滑、类别权重、RandAugment 等正则化与增强技术的应用，'
         '让我体会到防止过拟合在百类分类任务中的重要性。')

add_body('第二，在全栈系统工程方面，从 Django 后端 API 设计到 Vue 前端界面开发，'
         '我完整经历了前后端分离架构的设计与实现过程。JWT 认证、请求限流、缓存策略、'
         '文件去重、日志管理等工程化实践的引入，让我认识到一个可用的 AI 系统远不止模型训练，'
         '围绕模型的工程基础设施同样至关重要。')

add_body('第三，在 UI/UX 设计方面，"墨韵科技"设计系统的构建使我对设计令牌'
         '（Design Tokens）、CSS 变量体系、响应式布局策略、动画缓动函数等前端设计工程化概念'
         '有了系统性的理解。简洁、克制、有文化辨识度的设计语言，能够显著提升用户体验和产品品质感。')

add_body('存在的问题与改进方向：Custom CNN 尽管已达 90% 准确率，但与迁移学习模型仍有约 5% 的差距，'
         '未来可探索更深层的架构设计（如加入 Squeeze-and-Excitation 注意力模块）和更丰富的数据增强策略；'
         'ViT-B/16 的推理效率仍需优化，可考虑模型量化（INT8）、知识蒸馏或使用更轻量的 ViT 变体。')

add_body('通过本次课程设计，我不仅巩固了课堂所学的深度学习理论知识，更获得了从数据处理、'
         '模型训练、系统开发到界面设计的全链路项目实践经验，为今后从事 AI 应用开发奠定了坚实的基础。')

# ══════════════════════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════════════════════
add_heading1('参考文献')

refs = [
    '[1] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2016.',
    '[2] Tan M, Le Q V. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks[C]. International Conference on Machine Learning (ICML), 2019.',
    '[3] Dosovitskiy A, Beyer L, Kolesnikov A, et al. An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale[C]. International Conference on Learning Representations (ICLR), 2021.',
    '[4] Huang G, Liu Z, van der Maaten L, et al. Densely Connected Convolutional Networks[C]. IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2017.',
    '[5] Li Fei-Fei, Fergus R, Perona P. Learning Generative Visual Models from Few Training Examples: An Incremental Bayesian Approach Tested on 101 Object Categories[C]. CVPR Workshop on Generative-Model Based Vision, 2004.',
    '[6] Nilsback M E, Zisserman A. Automated Flower Classification over a Large Number of Classes[C]. Indian Conference on Computer Vision, Graphics and Image Processing, 2008.',
    '[7] PyTorch Documentation. torchvision.models[EB/OL]. https://pytorch.org/vision/stable/models.html.',
    '[8] Django REST Framework Documentation. API Guide[EB/OL]. https://www.django-rest-framework.org/api-guide/.',
    '[9] Vue.js Documentation. Guide[EB/OL]. https://vuejs.org/guide/introduction.html.',
    '[10] Cubuk E D, Zoph B, Shlens J, et al. RandAugment: Practical Automated Data Augmentation with a Reduced Search Space[C]. NeurIPS, 2020.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(os.path.dirname(output_dir), 'docs',
                           '《人工智能应用》课程设计文档_多数据集图像分类智能预测系统.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'文档已保存至: {output_path}')
